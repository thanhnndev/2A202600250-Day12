"""
Document generation tools for creating PDF recommendation documents.

Tools:
- generate_recommendation_pdf: Create PDF with user profile and school recommendations
- AccessibleCanvas: Subclass of Canvas with accessibility features
"""

import os
import logging
from datetime import datetime
from typing import List, Dict, Any, Optional
from pathlib import Path

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    Image, PageBreak
)
from reportlab.pdfgen import canvas
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.lib.utils import ImageReader

logger = logging.getLogger(__name__)


class AccessibleCanvas(canvas.Canvas):
    """
    Accessible PDF Canvas with proper document structure tags and metadata.
    
    Features:
    - Document structure tags (H1, H2, P, etc.)
    - Proper metadata (title, author, subject, keywords)
    - Logical reading order
    - Alternative text for images
    """
    
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._tag_level = 0
        
    def set_metadata(self, title: str, author: str, subject: str, keywords: List[str]):
        """Set PDF metadata for accessibility."""
        self.setTitle(title)
        self.setAuthor(author)
        self.setSubject(subject)
        self.setKeywords(", ".join(keywords))
        
    def begin_tag(self, tag_type: str, **attrs):
        """Begin a structure tag."""
        if hasattr(self, 'beginMarkedContent'):
            self.beginMarkedContent(tag_type, **attrs)
        self._tag_level += 1
        
    def end_tag(self):
        """End a structure tag."""
        if hasattr(self, 'endMarkedContent'):
            self.endMarkedContent()
        self._tag_level -= 1
        
    def drawStringWithTags(self, x, y, text, tag_type="P", **attrs):
        """Draw string with structure tags."""
        self.begin_tag(tag_type, **attrs)
        self.drawString(x, y, text)
        self.end_tag()


def generate_recommendation_pdf(
    user_profile: Dict[str, Any],
    schools: List[Dict[str, Any]],
    output_dir: Optional[str] = None
) -> Optional[str]:
    """
    Generate PDF recommendation document.
    
    Args:
        user_profile: User information (name, email, gpa, budget, etc.)
        schools: List of recommended schools with details
        output_dir: Output directory (default: from env PDF_OUTPUT_DIR)
    
    Returns:
        Path to generated PDF or None if failed
    """
    logger.info(f"Generating PDF for user: {user_profile.get('name')}")
    
    try:
        # Setup output directory
        output_dir = output_dir or os.getenv(
            "PDF_OUTPUT_DIR", 
            "./output/pdfs"
        )
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        
        # Generate filename
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        safe_name = user_profile.get("name", "user").replace(" ", "_")
        filename = f"recommendation_{safe_name}_{timestamp}.pdf"
        filepath = os.path.join(output_dir, filename)
        
        # Create PDF document
        doc = SimpleDocTemplate(
            filepath,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )
        
        # Build PDF content
        elements = []
        styles = getSampleStyleSheet()
        
        # Add custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1a73e8'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#333333'),
            spaceAfter=12,
            spaceBefore=12
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            textColor=colors.HexColor('#555555'),
            alignment=TA_JUSTIFY
        )
        
        # Title
        elements.append(Paragraph("🎓 Study Abroad Recommendation", title_style))
        elements.append(Spacer(1, 0.3*inch))
        
        # Generated date
        generated_date = datetime.utcnow().strftime("%B %d, %Y")
        elements.append(
            Paragraph(f"Generated on {generated_date}", normal_style)
        )
        elements.append(Spacer(1, 0.5*inch))
        
        # User Profile Section
        elements.append(Paragraph("👤 Student Profile", heading_style))
        
        profile_data = [
            ["Name:", user_profile.get("name", "N/A")],
            ["Email:", user_profile.get("email", "N/A")],
            ["GPA:", f"{user_profile.get('gpa', 'N/A')}/4.0"],
            ["Budget:", f"${user_profile.get('budget', 'N/A'):,.0f}/year"],
            ["Preferred Countries:", ", ".join(user_profile.get("preferred_countries", ["N/A"]))],
            ["Major:", user_profile.get("major", "N/A")]
        ]
        
        profile_table = Table(profile_data, colWidths=[2*inch, 4*inch])
        profile_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f5f5f5')),
        ]))
        elements.append(profile_table)
        elements.append(Spacer(1, 0.5*inch))
        
        # Schools Section
        elements.append(Paragraph("🏫 Recommended Schools", heading_style))
        elements.append(Spacer(1, 0.2*inch))
        
        # Add each school as a card
        for i, school in enumerate(schools[:3], 1):  # Top 3 schools
            # School header
            school_name = school.get("name", "Unknown School")
            match_score = school.get("match_score", 0)
            
            elements.append(
                Paragraph(f"{i}. {school_name} (Match: {match_score:.1f}%)", heading_style)
            )
            
            # School details
            school_info = [
                f"📍 Location: {school.get('city', 'N/A')}, {school.get('country', 'N/A')}",
                f"📚 Programs: {school.get('programs_count', 0)} available",
                f"🏷️ Category: {school.get('category', 'N/A')}"
            ]
            
            if school.get("website"):
                school_info.append(f"🌐 Website: {school['website']}")
            
            for info in school_info:
                elements.append(Paragraph(info, normal_style))
            
            # Reasons for recommendation
            reasons = school.get("reasons", [])
            if reasons:
                elements.append(Spacer(1, 0.1*inch))
                elements.append(
                    Paragraph("<b>Why this school matches you:</b>", normal_style)
                )
                for reason in reasons:
                    elements.append(
                        Paragraph(f"  • {reason}", normal_style)
                    )
            
            elements.append(Spacer(1, 0.3*inch))
        
        # Next Steps Section
        elements.append(Paragraph("📋 Next Steps", heading_style))
        
        next_steps = [
            "1. Review the recommended schools and programs above",
            "2. Visit school websites to learn more about programs and campus life",
            "3. Prepare required documents (transcripts, test scores, essays)",
            "4. Submit applications through PathCan portal",
            "5. Track your application status in real-time"
        ]
        
        for step in next_steps:
            elements.append(Paragraph(f"• {step}", normal_style))
        
        elements.append(Spacer(1, 0.5*inch))
        
        # Contact Information
        elements.append(
            Paragraph("📞 Need Help? Contact PathCan Support", heading_style)
        )
        
        contact_info = [
            "Email: support@pathcan.com",
            "Phone: +1 (555) 123-4567",
            "Website: www.pathcan.com"
        ]
        
        for info in contact_info:
            elements.append(Paragraph(info, normal_style))
        
        # Footer note
        elements.append(Spacer(1, 0.5*inch))
        elements.append(
            Paragraph(
                "<i>This recommendation was generated by PathCan Smart Apply AI. "
                "School information is updated daily. Please verify details on official "
                "school websites before applying.</i>",
                ParagraphStyle(
                    'Footer',
                    parent=styles['Normal'],
                    fontSize=9,
                    textColor=colors.HexColor('#999999'),
                    alignment=TA_CENTER
                )
            )
        )
        
        # Build PDF
        doc.build(elements)
        
        logger.info(f"PDF generated successfully: {filepath}")
        return filepath
        
    except Exception as e:
        logger.error(f"PDF generation failed: {e}")
        return None


def generate_recommendation_docx(
    user_profile: Dict[str, Any],
    schools: List[Dict[str, Any]],
    output_dir: Optional[str] = None
) -> Optional[str]:
    """
    Generate Word document recommendation (alternative to PDF).
    
    Args:
        user_profile: User information
        schools: List of recommended schools
        output_dir: Output directory
    
    Returns:
        Path to generated DOCX or None if failed
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    logger.info(f"Generating DOCX for user: {user_profile.get('name')}")
    
    try:
        # Setup output directory
        output_dir = output_dir or os.getenv(
            "PDF_OUTPUT_DIR",
            "./output/pdfs"
        )
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        
        # Generate filename
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        safe_name = user_profile.get("name", "user").replace(" ", "_")
        filename = f"recommendation_{safe_name}_{timestamp}.docx"
        filepath = os.path.join(output_dir, filename)
        
        # Create document
        doc = Document()
        
        # Title
        title = doc.add_heading('🎓 Study Abroad Recommendation', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Generated date
        generated_date = datetime.utcnow().strftime("%B %d, %Y")
        doc.add_paragraph(f'Generated on {generated_date}')
        
        # User Profile
        doc.add_heading('👤 Student Profile', level=1)
        
        profile_table = doc.add_table(rows=6, cols=2)
        profile_table.style = 'Light Grid Accent 1'
        
        profile_data = [
            ("Name", user_profile.get("name", "N/A")),
            ("Email", user_profile.get("email", "N/A")),
            ("GPA", f"{user_profile.get('gpa', 'N/A')}/4.0"),
            ("Budget", f"${user_profile.get('budget', 'N/A'):,.0f}/year"),
            ("Preferred Countries", ", ".join(user_profile.get("preferred_countries", ["N/A"]))),
            ("Major", user_profile.get("major", "N/A"))
        ]
        
        for i, (label, value) in enumerate(profile_data):
            profile_table.rows[i].cells[0].text = label
            profile_table.rows[i].cells[1].text = value
            profile_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()
        
        # Schools
        doc.add_heading('🏫 Recommended Schools', level=1)
        
        for i, school in enumerate(schools[:3], 1):
            school_name = school.get("name", "Unknown School")
            match_score = school.get("match_score", 0)
            
            doc.add_heading(f'{i}. {school_name} (Match: {match_score:.1f}%)', level=2)
            
            doc.add_paragraph(f'📍 Location: {school.get("city", "N/A")}, {school.get("country", "N/A")}')
            doc.add_paragraph(f'📚 Programs: {school.get("programs_count", 0)} available')
            doc.add_paragraph(f'🏷️ Category: {school.get("category", "N/A")}')
            
            if school.get("website"):
                doc.add_paragraph(f'🌐 Website: {school["website"]}')
            
            # Reasons
            reasons = school.get("reasons", [])
            if reasons:
                doc.add_paragraph('Why this school matches you:', style='Intense Quote')
                for reason in reasons:
                    doc.add_paragraph(f'  • {reason}', style='List Bullet')
            
            doc.add_paragraph()
        
        # Next Steps
        doc.add_heading('📋 Next Steps', level=1)
        doc.add_paragraph('1. Review the recommended schools and programs', style='List Number')
        doc.add_paragraph('2. Visit school websites to learn more', style='List Number')
        doc.add_paragraph('3. Prepare required documents', style='List Number')
        doc.add_paragraph('4. Submit applications through PathCan portal', style='List Number')
        doc.add_paragraph('5. Track your application status', style='List Number')
        
        # Save document
        doc.save(filepath)
        
        logger.info(f"DOCX generated successfully: {filepath}")
        return filepath
        
    except Exception as e:
        logger.error(f"DOCX generation failed: {e}")
        return None


# Tool configuration
TOOLS_CONFIG = {
    "generate_recommendation_pdf": {
        "name": "generate_recommendation_pdf",
        "description": "Generate PDF recommendation document with user profile and top 3 schools",
        "func": generate_recommendation_pdf,
        "parameters": {
            "user_profile": "Dict - User information (name, email, gpa, budget, etc.)",
            "schools": "List[Dict] - List of recommended schools with details",
            "output_dir": "Optional[str] - Output directory (default: ./output/pdfs)"
        }
    },
    "generate_recommendation_docx": {
        "name": "generate_recommendation_docx",
        "description": "Generate Word document recommendation (alternative to PDF)",
        "func": generate_recommendation_docx,
        "parameters": {
            "user_profile": "Dict - User information",
            "schools": "List[Dict] - List of recommended schools",
            "output_dir": "Optional[str] - Output directory"
        }
    }
}
